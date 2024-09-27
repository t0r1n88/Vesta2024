"""
Функции для создания документов из шаблонов
"""

import pandas as pd
import numpy as np
import os
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from tkinter import messagebox
from jinja2 import exceptions
import time
import datetime
import warnings

warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')
warnings.simplefilter(action='ignore', category=DeprecationWarning)
warnings.simplefilter(action='ignore', category=UserWarning)
pd.options.mode.chained_assignment = None
import platform
import logging
import tempfile
import re

logging.basicConfig(
    level=logging.WARNING,
    filename="error.log",
    filemode='w',
    # чтобы файл лога перезаписывался  при каждом запуске.Чтобы избежать больших простыней. По умолчанию идет 'a'
    format="%(asctime)s - %(module)s - %(levelname)s - %(funcName)s: %(lineno)d - %(message)s",
    datefmt='%H:%M:%S',
)
class CheckBoxException(Exception):
    """
    Класс для вызовы исключения в случае если неправильно выставлены чекбоксы
    """
    pass


class NotFoundValue(Exception):
    """
    Класс для обозначения того что значение не найдено
    """
    pass

class NotNumberColumn(Exception):
    """
    Исключение для обработки варианта когда в таблице нет колонки с таким порядковым номером
    """
    pass

class NoMoreNumberColumn(Exception):
    """
    Исключение для обработки варианта если введено больше 2 порядковых номеров колонок для создания структуры папок

    """
    pass


def create_doc_convert_date(cell):
    """
    Функция для конвертации даты при создании документов
    :param cell:
    :return:
    """
    try:
        string_date = datetime.datetime.strftime(cell, '%d.%m.%Y')
        return string_date
    except ValueError:
        return 'Не удалось конвертировать дату.Проверьте значение ячейки!!!'
    except TypeError:
        return 'Не удалось конвертировать дату.Проверьте значение ячейки!!!'


def processing_date_column(df, lst_columns):
    """
    Функция для обработки столбцов с датами. конвертация в строку формата ДД.ММ.ГГГГ
    """
    # получаем первую строку
    first_row = df.iloc[0, lst_columns]

    lst_first_row = list(first_row)  # Превращаем строку в список
    lst_date_columns = []  # Создаем список куда будем сохранять колонки в которых находятся даты
    tupl_row = list(zip(lst_columns,
                        lst_first_row))  # Создаем список кортежей формата (номер колонки,значение строки в этой колонке)

    for idx, value in tupl_row:  # Перебираем кортеж
        result = check_date_columns(idx, value)  # проверяем является ли значение датой
        if result:  # если да то добавляем список порядковый номер колонки
            lst_date_columns.append(result)
        else:  # иначе проверяем следующее значение
            continue
    for i in lst_date_columns:  # Перебираем список с колонками дат, превращаем их в даты и конвертируем в нужный строковый формат
        df.iloc[:, i] = pd.to_datetime(df.iloc[:, i], errors='coerce', dayfirst=True)
        df.iloc[:, i] = df.iloc[:, i].apply(create_doc_convert_date)

def check_date_columns(i, value):
    """
    Функция для проверки типа колонки. Необходимо найти колонки с датой
    :param i:
    :param value:
    :return:
    """
    try:
        itog = pd.to_datetime(str(value), infer_datetime_format=True)
    except:
        pass
    else:
        return i

def clean_value(value):
    """
    Функция для обработки значений колонки от  пустых пробелов,нан
    :param value: значение ячейки
    :return: очищенное значение
    """
    if value is np.nan:
        return 'Не заполнено'
    str_value = str(value)
    if str_value == '':
        return 'Не заполнено'
    elif str_value ==' ':
        return 'Не заполнено'

    return str_value



def combine_all_docx(filename_master, files_lst,mode_pdf,path_to_end_folder_doc,name_os):
    """
    Функция для объединения файлов Word взято отсюда
    https://stackoverflow.com/questions/24872527/combine-word-document-using-python-docx
    :param filename_master: базовый файл
    :param files_list: список с созданными файлами
    :return: итоговый файл
    """

    # Получаем текущее время
    t = time.localtime()
    current_time = time.strftime('%H_%M_%S', t)

    number_of_sections = len(files_lst)
    # Открываем и обрабатываем базовый файл
    master = Document(filename_master)
    composer = Composer(master)
    # Перебираем и добавляем файлы к базовому
    for i in range(0, number_of_sections):
        doc_temp = Document(files_lst[i])
        composer.append(doc_temp)
    # Сохраняем файл
    composer.save(f"{path_to_end_folder_doc}/Объединенный файл от {current_time}.docx")
    if mode_pdf == 'Yes':
        if name_os == 'Windows':
            convert(f"{path_to_end_folder_doc}/Объединенный файл от {current_time}.docx",
                f"{path_to_end_folder_doc}/Объединенный файл от {current_time}.pdf", keep_active=True)
        else:
            raise NotImplementedError

def prepare_entry_str(raw_str:str,pattern:str,repl_str:str,sep_lst:str)->list:
    """
    Функция для очистки строки от лишних символов и уменьшения на единицу (для нумерации с нуля)
    :param raw_str: обрабатываемая строка
    :param pattern: паттерн для замены символов
    :param repl_str: строка на которую нужно заменять символы
    :param sep_lst: разделитель по которому будет делиться список
    :return: список
    """
    raw_str = str(raw_str).replace('.',',')
    number_column_folder_structure = re.sub(pattern,repl_str,raw_str) # убираем из строки все лишние символы
    lst_number_column_folder_structure = number_column_folder_structure.split(sep_lst) # создаем список по запятой
    # отбрасываем возможные лишние элементы из за лишних запятых
    lst_number_column_folder_structure = [value for value in lst_number_column_folder_structure if value]
    # заменяем 0 на единицу
    lst_number_column_folder_structure = ['1' if x == '0' else x for x in lst_number_column_folder_structure]
    # Превращаем в числа и отнимаем 1 чтобы соответствовать индексам питона
    lst_number_column_folder_structure = list(map(lambda x:int(x)-1,lst_number_column_folder_structure))
    return lst_number_column_folder_structure

def save_result_file(finish_path:str,name_file:str,doc:DocxTemplate,idx:int,mode_pdf:str,name_os):
    """
    Функция для сохранения результатов
    :param finish_path: путь к папке сохранения
    :param name_file: название файла
    :param doc: объект DocxTemplate
    :param idx: счетчик
    :param mode_pdf: чекбокс сохранения PDF
    :param name_os: операционная система
    :return:
    """
    if os.path.exists(f'{finish_path}/{name_file}.docx'):
        doc.save(f'{finish_path}/{name_file}_{idx}.docx')
        if mode_pdf == 'Yes':
            if name_os == 'Windows':
                if not os.path.exists(f'{finish_path}/PDF'):
                    os.makedirs(f'{finish_path}/PDF')
                convert(f'{finish_path}/{name_file}_{idx}.docx', f'{finish_path}/PDF/{name_file}_{idx}.pdf',
                        keep_active=True)
            else:
                raise NotImplementedError
    else:
        doc.save(f'{finish_path}/{name_file}.docx')
        if mode_pdf == 'Yes':
            if name_os == 'Windows':
                if not os.path.exists(f'{finish_path}/PDF'):
                    os.makedirs(f'{finish_path}/PDF')
                convert(f'{finish_path}/{name_file}.docx', f'{finish_path}/PDF/{name_file}.pdf',
                        keep_active=True)
            else:
                raise NotImplementedError

def generate_docs_from_template(name_file_template_doc, name_file_data_doc,name_column, name_type_file,path_to_end_folder_doc, name_value_column, mode_pdf,
                                mode_combine, mode_group,mode_structure_folder,number_structure_folder):
    """
    Функция для создания однотипных документов из шаблона Word и списка Excel
    :param name_file_template_doc:путь к файлу шаблону на основе которого будут генерироваться документы
    :param name_file_data_doc: путь к файлу Excel с данными которые подставляются в шаблон
    :param name_column: название колонки в таблице данные из которой будут использоватьс для создания названий документов
    :param name_type_file: название создаваемых документов например Согласие,Справка и т.д.
    :param path_to_end_folder_doc: путь к папке куда будут сохраняться файлы
    :param name_value_column: Значение из колонки name_type_file по которому будет создан единичный документ
    :param mode_pdf: чекбокс отвечающий за режим работы с pdf если Yes то будет создавать дополнительно pdf документ
    :param mode_combine:чекбокс отвечающий за режим объединения файлов. Если Yes то все документы будут объединены в один
    файл, если No то будет создаваться отдельный документ на каждую строчку исходной таблицы
    :param mode_group: чекбокс отвечающий за режим создания отдельного файла. Если Yes то можно создать один файл по значению
     из колонки name_value_column
    :param mode_structure_folder: чекбокс отвечающий за создание структуры папок
    :param number_structure_folder: строка с порядковыми номерами колонок разделенными запятой
    :return: Создает в зависимости от выбранного режима файлы Word из шаблона
    """
    try:
        name_os = platform.system() # получаем платформу на которой запущена программа
        # Считываем данные
        # Добавил параметр dtype =str чтобы данные не преобразовались а использовались так как в таблице
        df = pd.read_excel(name_file_data_doc, dtype=str)
        df[name_column] = df[name_column].apply(clean_value) # преобразовываем колонку меняя пустые значения и пустые пробелы на Не заполнено
        used_name_file = set()  # множество для уже использованных имен файлов
        # Заполняем Nan
        df.fillna(' ', inplace=True)
        lst_date_columns = []

        for idx, column in enumerate(df.columns):
            if 'дата' in column.lower():
                lst_date_columns.append(idx)

        # Конвертируем в пригодный строковый формат
        for i in lst_date_columns:
            df.iloc[:, i] = pd.to_datetime(df.iloc[:, i], errors='coerce', dayfirst=True)
            df.iloc[:, i] = df.iloc[:, i].apply(create_doc_convert_date)

        # Конвертируем датафрейм в список словарей
        data = df.to_dict('records')
        if mode_structure_folder == 'No':
            # В зависимости от состояния чекбоксов обрабатываем файлы
            if mode_combine == 'No':
                if mode_group == 'No':
                    # Создаем в цикле документы
                    for idx, row in enumerate(data):
                        doc = DocxTemplate(name_file_template_doc)
                        context = row
                        # print(context)
                        doc.render(context)
                        # Сохраняенм файл
                        # получаем название файла и убираем недопустимые символы < > : " /\ | ? *
                        name_file = row[name_column]
                        name_file = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_', name_file)
                        # проверяем файл на наличие, если файл с таким названием уже существует то добавляем окончание
                        if name_file in used_name_file:
                            name_file = f'{name_file}_{idx}'

                        doc.save(f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.docx')
                        used_name_file.add(name_file)
                        if mode_pdf == 'Yes':
                            if name_os == 'Windows':
                                convert(f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.docx',
                                    f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.pdf', keep_active=True)
                            else:
                                raise NotImplementedError
                else:
                    # Отбираем по значению строку

                    single_df = df[df[name_column] == name_value_column]
                    # Конвертируем датафрейм в список словарей
                    single_data = single_df.to_dict('records')
                    # Проверяем количество найденных совпадений
                    # очищаем от запрещенных символов
                    name_file = name_value_column
                    name_file = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_', name_file)
                    if len(single_data) == 1:
                        for row in single_data:
                            doc = DocxTemplate(name_file_template_doc)
                            doc.render(row)
                            # Сохраняенм файл
                            doc.save(f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.docx')
                            if mode_pdf == 'Yes':
                                if name_os == 'Windows':
                                    convert(f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.docx',
                                        f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.pdf', keep_active=True)
                                else:
                                    raise NotImplementedError

                    elif len(single_data) > 1:
                        for idx, row in enumerate(single_data):
                            doc = DocxTemplate(name_file_template_doc)
                            doc.render(row)
                            # Сохраняем файл

                            doc.save(f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}_{idx}.docx')
                            if mode_pdf == 'Yes':
                                if name_os == 'Windows':
                                    convert(f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}_{idx}.docx',
                                        f'{path_to_end_folder_doc}/{name_type_file} {name_file[:80]}.pdf', keep_active=True)
                                else:
                                    raise NotImplementedError
                    else:
                        raise NotFoundValue



            else:
                if mode_group == 'No':
                    # Список с созданными файлами
                    files_lst = []

                    # Добавляем разрыв в шаблон
                    # Открываем шаблон
                    doc_page_break = Document(name_file_template_doc)
                    # Добавляем разрыв страницы
                    doc_page_break.add_page_break()
                    template_page_break_path = os.path.dirname(name_file_template_doc)
                    # Сохраняем изменения в новом файле
                    doc_page_break.save(f'{template_page_break_path}/page_break.docx')
                    # Создаем временную папку
                    with tempfile.TemporaryDirectory() as tmpdirname:
                        print('created temporary directory', tmpdirname)
                        # Создаем и сохраняем во временную папку созданные документы Word
                        for idx,row in enumerate(data):
                            # Открываем файл
                            doc = DocxTemplate(f'{template_page_break_path}/page_break.docx')
                            context = row
                            doc.render(context)
                            # Сохраняем файл
                            # очищаем от запрещенных символов
                            name_file = f'{row[name_column]}'
                            name_file = re.sub(r'[\r\b\n\t<> :"?*|\\/]', '_', name_file)

                            doc.save(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                            # Добавляем путь к файлу в список
                            files_lst.append(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                        # Получаем базовый файл
                        main_doc = files_lst.pop(0)
                        # Запускаем функцию
                        combine_all_docx(main_doc, files_lst,mode_pdf,path_to_end_folder_doc,name_os)
                        # Удаляем файл с разрывом страницы
                        try:
                            os.remove(f'{template_page_break_path}/page_break.docx')
                        except OSError as e:
                            print("Ошибка при попытке удаления файла: {}".format(e))
                else:
                    raise CheckBoxException
        else:
            # очищаем строку от лишних символов и превращаем в список номеров колонок
            lst_number_column_folder_structure = prepare_entry_str(number_structure_folder, r'[^\d,]', '', ',')
            # проверяем длину списка не более 2 и не равно 0
            if len(lst_number_column_folder_structure) == 0 or len(lst_number_column_folder_structure) > 3:
                raise NoMoreNumberColumn
            # проверяем чтобы номер колонки не превышал количество колонок в датафрейме
            for number_column in lst_number_column_folder_structure:
                if number_column > df.shape[1]-1:
                    raise NotNumberColumn

            if len(lst_number_column_folder_structure) == 1:
                # Если нужно создавать одноуровневую структуру
                # получаем название колонки
                main_layer_name_column = df.columns[lst_number_column_folder_structure[0]]
                # Заменяем пробелы на Не заполнено
                df[main_layer_name_column] = df[main_layer_name_column].apply(lambda x:'Не заполнено' if x == ' ' else x)
                lst_unique_value = df[main_layer_name_column].unique()  # получаем список уникальных значений
                for name_folder in lst_unique_value:
                    temp_df = df[df[main_layer_name_column] == name_folder]  # фильтруем по названию
                    # Создаем название для папки
                    clean_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                               name_folder)  # очищаем название от лишних символов
                    finish_path = f'{path_to_end_folder_doc}/{clean_name_folder}'
                    if not os.path.exists(finish_path):
                        os.makedirs(finish_path)
                    temp_df = temp_df.applymap(
                        lambda x: str.replace(x, 'Не заполнено', '') if isinstance(x, str) else x)
                    data = temp_df.to_dict('records') # делаем из датафрейма список словарей

                    if mode_combine == 'No':
                        if mode_group == 'No':
                            # Создаем в цикле документы
                            for idx, row in enumerate(data):
                                doc = DocxTemplate(name_file_template_doc)
                                context = row
                                doc.render(context)
                                name_file = f'{name_type_file} {row[name_column]}'
                                name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                threshold_name = 200 - (len(finish_path) + 10)
                                if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                    raise OSError
                                name_file = name_file[:threshold_name]  # ограничиваем название файла
                                # Сохраняем файл
                                save_result_file(finish_path, name_file, doc, idx, mode_pdf,name_os)
                        else:
                            raise CheckBoxException
                    else:
                        if mode_group == 'No':
                            # Список с созданными файлами
                            files_lst = []

                            # Добавляем разрыв в шаблон
                            # Открываем шаблон
                            doc_page_break = Document(name_file_template_doc)
                            # Добавляем разрыв страницы
                            doc_page_break.add_page_break()
                            template_page_break_path = os.path.dirname(name_file_template_doc)
                            # Сохраняем изменения в новом файле
                            doc_page_break.save(f'{template_page_break_path}/page_break.docx')
                            # Создаем временную папку
                            with tempfile.TemporaryDirectory() as tmpdirname:
                                print('created temporary directory', tmpdirname)
                                # Создаем и сохраняем во временную папку созданные документы Word
                                for idx, row in enumerate(data):
                                    # Открываем файл
                                    doc = DocxTemplate(f'{template_page_break_path}/page_break.docx')
                                    context = row
                                    doc.render(context)
                                    # Сохраняем файл
                                    # очищаем от запрещенных символов
                                    name_file = f'{row[name_column]}'
                                    name_file = re.sub(r'[\r\b\n\t<> :"?*|\\/]', '_', name_file)

                                    doc.save(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                                    # Добавляем путь к файлу в список
                                    files_lst.append(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                                # Получаем базовый файл
                                main_doc = files_lst.pop(0)
                                # Запускаем функцию
                                combine_all_docx(main_doc, files_lst, mode_pdf, finish_path, name_os)
                                # Удаляем файл с разрывом страницы
                                try:
                                    os.remove(f'{template_page_break_path}/page_break.docx')
                                except OSError as e:
                                    print("Ошибка при попытке удаления файла: {}".format(e))
                        else:
                            raise CheckBoxException

            if len(lst_number_column_folder_structure) == 2:
                # Создаем папки для двухзначной структуры
                name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]

                # Заменяем пробелы на Не заполнено
                df[name_first_layer_column] = df[name_first_layer_column].apply(lambda x:'Не заполнено' if x == ' ' else x)
                df[name_second_layer_column] = df[name_second_layer_column].apply(lambda x:'Не заполнено' if x == ' ' else x)

                lst_unique_value_first_layer = df[
                    name_first_layer_column].unique()  # получаем список уникальных значений
                for first_name_folder in lst_unique_value_first_layer:
                    clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                     first_name_folder)  # очищаем название от лишних символов

                    # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                    temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                    lst_unique_value_second_layer = temp_df_first_layer[
                        name_second_layer_column].unique()  # получаем список уникальных значений
                    # фильтруем по значениям колонки второго уровня
                    for second_name_folder in lst_unique_value_second_layer:
                        temp_df_second_layer = temp_df_first_layer[
                            temp_df_first_layer[name_second_layer_column] == second_name_folder]
                        clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                          second_name_folder)  # очищаем название от лишних символов

                        finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}'
                        if not os.path.exists(finish_path):
                            os.makedirs(finish_path)
                        temp_df_second_layer = temp_df_second_layer.applymap(
                            lambda x: str.replace(x, 'Не заполнено', '') if isinstance(x, str) else x)
                        data = temp_df_second_layer.to_dict('records')  # конвертируем в список словарей
                        if mode_combine == 'No':
                            if mode_group == 'No':
                                # Создаем в цикле документы
                                for idx, row in enumerate(data):
                                    doc = DocxTemplate(name_file_template_doc)
                                    context = row
                                    doc.render(context)
                                    name_file = f'{name_type_file} {row[name_column]}'
                                    name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                    threshold_name = 200 - (len(finish_path) + 10)
                                    if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                        raise OSError
                                    name_file = name_file[:threshold_name]  # ограничиваем название файла
                                    # Сохраняем файл
                                    save_result_file(finish_path, name_file, doc, idx, mode_pdf, name_os)
                            else:
                                raise CheckBoxException
                        else:
                            if mode_group == 'No':
                                # Список с созданными файлами
                                files_lst = []

                                # Добавляем разрыв в шаблон
                                # Открываем шаблон
                                doc_page_break = Document(name_file_template_doc)
                                # Добавляем разрыв страницы
                                doc_page_break.add_page_break()
                                template_page_break_path = os.path.dirname(name_file_template_doc)
                                # Сохраняем изменения в новом файле
                                doc_page_break.save(f'{template_page_break_path}/page_break.docx')
                                # Создаем временную папку
                                with tempfile.TemporaryDirectory() as tmpdirname:
                                    print('created temporary directory', tmpdirname)
                                    # Создаем и сохраняем во временную папку созданные документы Word
                                    for idx, row in enumerate(data):
                                        # Открываем файл
                                        doc = DocxTemplate(f'{template_page_break_path}/page_break.docx')
                                        context = row
                                        doc.render(context)
                                        # Сохраняем файл
                                        # очищаем от запрещенных символов
                                        name_file = f'{row[name_column]}'
                                        name_file = re.sub(r'[\r\b\n\t<> :"?*|\\/]', '_', name_file)

                                        doc.save(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                                        # Добавляем путь к файлу в список
                                        files_lst.append(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                                    # Получаем базовый файл
                                    main_doc = files_lst.pop(0)
                                    # Запускаем функцию
                                    combine_all_docx(main_doc, files_lst, mode_pdf, finish_path, name_os)
                                    # Удаляем файл с разрывом страницы
                                    try:
                                        os.remove(f'{template_page_break_path}/page_break.docx')
                                    except OSError as e:
                                        print("Ошибка при попытке удаления файла: {}".format(e))
                            else:
                                raise CheckBoxException
            if len(lst_number_column_folder_structure) == 3:
                # Если нужно создавать трехуровневую структуру Например Школа-Класс-буква класса
                # получаем названия колонок для трех уровней
                name_first_layer_column = df.columns[lst_number_column_folder_structure[0]]
                name_second_layer_column = df.columns[lst_number_column_folder_structure[1]]
                name_third_layer_column = df.columns[lst_number_column_folder_structure[2]]

                # Заменяем пробелы на Не заполнено
                df[name_first_layer_column] = df[name_first_layer_column].apply(lambda x:'Не заполнено' if x == ' ' else x)
                df[name_second_layer_column] = df[name_second_layer_column].apply(lambda x:'Не заполнено' if x == ' ' else x)
                df[name_third_layer_column] = df[name_third_layer_column].apply(lambda x:'Не заполнено' if x == ' ' else x)

                lst_unique_value_first_layer = df[
                    name_first_layer_column].unique()  # получаем список уникальных значений
                for first_name_folder in lst_unique_value_first_layer:
                    clean_first_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                     first_name_folder)  # очищаем название от лишних символов

                    # получаем отфильтрованный датафрейм по значениям колонки первого уровня
                    temp_df_first_layer = df[df[name_first_layer_column] == first_name_folder]  # фильтруем по названию
                    lst_unique_value_second_layer = temp_df_first_layer[
                        name_second_layer_column].unique()  # получаем список уникальных значений второго уровня
                    # фильтруем по значениям колонки второго уровня
                    for second_name_folder in lst_unique_value_second_layer:
                        temp_df_second_layer = temp_df_first_layer[
                            temp_df_first_layer[name_second_layer_column] == second_name_folder]
                        clean_second_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                          second_name_folder)  # очищаем название от лишних символов
                        lst_unique_value_third_layer = temp_df_second_layer[
                            name_third_layer_column].unique()  # получаем список уникальных значений третьего уровня
                        for third_name_folder in lst_unique_value_third_layer:
                            clean_third_name_folder = re.sub(r'[\r\b\n\t<>:"?*|\\/]', '_',
                                                             third_name_folder)  # очищаем название от лишних символов
                            temp_df_third_layer = temp_df_second_layer[
                                temp_df_second_layer[name_third_layer_column] == third_name_folder]



                            finish_path = f'{path_to_end_folder_doc}/{clean_first_name_folder}/{clean_second_name_folder}/{clean_third_name_folder}'
                            if not os.path.exists(finish_path):
                                os.makedirs(finish_path)
                            # заменяем перед записью документа Не заполнено на пробел
                            temp_df_third_layer = temp_df_third_layer.applymap(lambda x: str.replace(x,'Не заполнено','') if isinstance(x,str) else x)
                            data = temp_df_third_layer.to_dict('records')  # конвертируем в список словарей
                            if mode_combine == 'No':
                                if mode_group == 'No':
                                    # Создаем в цикле документы
                                    for idx, row in enumerate(data):
                                        doc = DocxTemplate(name_file_template_doc)
                                        context = row
                                        doc.render(context)
                                        name_file = f'{name_type_file} {row[name_column]}'
                                        name_file = re.sub(r'[<> :"?*|\\/]', ' ', name_file)
                                        threshold_name = 200 - (len(finish_path) + 10)
                                        if threshold_name <= 0:  # если путь к папке слишком длинный вызываем исключение
                                            raise OSError
                                        name_file = name_file[:threshold_name]  # ограничиваем название файла
                                        # Сохраняем файл
                                        save_result_file(finish_path, name_file, doc, idx, mode_pdf, name_os)
                                else:
                                    raise CheckBoxException
                            else:
                                if mode_group == 'No':
                                    # Список с созданными файлами
                                    files_lst = []

                                    # Добавляем разрыв в шаблон
                                    # Открываем шаблон
                                    doc_page_break = Document(name_file_template_doc)
                                    # Добавляем разрыв страницы
                                    doc_page_break.add_page_break()
                                    template_page_break_path = os.path.dirname(name_file_template_doc)
                                    # Сохраняем изменения в новом файле
                                    doc_page_break.save(f'{template_page_break_path}/page_break.docx')
                                    # Создаем временную папку
                                    with tempfile.TemporaryDirectory() as tmpdirname:
                                        print('created temporary directory', tmpdirname)
                                        # Создаем и сохраняем во временную папку созданные документы Word
                                        for idx, row in enumerate(data):
                                            # Открываем файл
                                            doc = DocxTemplate(f'{template_page_break_path}/page_break.docx')
                                            context = row
                                            doc.render(context)
                                            # Сохраняем файл
                                            # очищаем от запрещенных символов
                                            name_file = f'{row[name_column]}'
                                            name_file = re.sub(r'[\r\b\n\t<> :"?*|\\/]', '_', name_file)

                                            doc.save(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                                            # Добавляем путь к файлу в список
                                            files_lst.append(f'{tmpdirname}/{name_file[:80]}_{idx}.docx')
                                        # Получаем базовый файл
                                        main_doc = files_lst.pop(0)
                                        # Запускаем функцию
                                        combine_all_docx(main_doc, files_lst, mode_pdf, finish_path, name_os)
                                        # Удаляем файл с разрывом страницы
                                        try:
                                            os.remove(f'{template_page_break_path}/page_break.docx')
                                        except OSError as e:
                                            print("Ошибка при попытке удаления файла: {}".format(e))
                                else:
                                    raise CheckBoxException

    except NameError as e:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Выберите шаблон,файл с данными и папку куда будут генерироваться файлы')
        logging.exception('AN ERROR HAS OCCURRED')
    except KeyError as e:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'В таблице не найдена указанная колонка {e.args}')
    except PermissionError:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Закройте все файлы Word созданные Вестой')
        logging.exception('AN ERROR HAS OCCURRED')
    except FileNotFoundError:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Перенесите файлы, конечную папку с которой вы работете в корень диска. Проблема может быть\n '
                             f'в слишком длинном пути к обрабатываемым файлам или конечной папке.')
    except exceptions.TemplateSyntaxError:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Ошибка в оформлении вставляемых значений в шаблоне\n'
                             f'Проверьте свой шаблон на наличие следующих ошибок:\n'
                             f'1) Вставляемые значения должны быть оформлены двойными фигурными скобками\n'
                             f'{{{{Вставляемое_значение}}}}\n'
                             f'2) В названии колонки в таблице откуда берутся данные - есть пробелы,цифры,знаки пунктуации и т.п.\n'
                             f'в названии колонки должны быть только буквы и нижнее подчеркивание.\n'
                             f'{{{{Дата_рождения}}}}')

    except NotImplementedError as e:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Создание pdf файлов работает ТОЛЬКО в WIndows, уберите галочку из чекбокса создания pdf ')
    except CheckBoxException:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Уберите галочку из чекбокса Поставьте галочку, если вам нужно создать один документ\nдля конкретного значения (например для определенного ФИО)'
                             )
    except NotFoundValue:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Указанное значение не найдено в выбранной колонке\nПроверьте наличие такого значения в таблице'
                             )
    except NotNumberColumn:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Колонки с таким порядковым номером нет в таблице.\nПроверьте правильность введенных данных.'
                             )
    except NoMoreNumberColumn:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Проверьте количество введенных порядковых номеров колонок.\n'
                             f'Не более 3 чисел разделенных запятыми.\n'
                             f'Например 3,12,8'
                             )
    except OSError:
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             f'Слишком длинный путь к файлу. Выберите конечную папку в корне диска или выберите другие колонки для создания структуры папок')
    except:
        logging.exception('AN ERROR HAS OCCURRED')
        messagebox.showerror('Веста Обработка таблиц и создание документов',
                             'Возникла ошибка!!! Подробности ошибки в файле error.log')

    else:
        messagebox.showinfo('Веста Обработка таблиц и создание документов', 'Создание документов завершено!')

if __name__ == '__main__':
    name_column_main = 'ФИО'
    name_type_file_main = 'Справка'
    name_value_column_main = 'Алехин Данила Прокопьевич'
    mode_pdf_main = 'Yes'
    name_file_template_doc_main = 'data/Создание документов/Пример Шаблон согласия.docx'
    name_file_data_doc_main = 'data/Создание документов/Таблица для заполнения согласия.xlsx'
    path_to_end_folder_doc_main = 'data/result'
    mode_combine_main = 'No'
    mode_group_main = 'No'
    main_mode_structure_folder = 'Yes'
    main_structure_folder = '10'

    generate_docs_from_template(name_file_template_doc_main,name_file_data_doc_main,name_column_main, name_type_file_main, path_to_end_folder_doc_main,
                                name_value_column_main, mode_pdf_main,
                                mode_combine_main, mode_group_main,main_mode_structure_folder,main_structure_folder)

    print('Lindy Booth')